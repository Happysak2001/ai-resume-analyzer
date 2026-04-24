import os
import io
import re
import html
import streamlit as st
from groq import Groq
import pdfplumber
import docx2txt

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.units import inch


st.set_page_config(page_title="AI Resume Rewriter", layout="wide")

st.title("AI Resume Rewriter")
st.write("Upload your resume and paste the job description below.")

uploaded_file = st.file_uploader("Upload Resume (PDF or DOCX)", type=["pdf", "docx"])
job_description = st.text_area("Paste Job Description", height=220)

resume_text = ""

if uploaded_file is not None:
    if uploaded_file.name.endswith(".pdf"):
        with pdfplumber.open(uploaded_file) as pdf:
            for page in pdf.pages:
                resume_text += page.extract_text() or ""

    elif uploaded_file.name.endswith(".docx"):
        resume_text = docx2txt.process(uploaded_file)

    st.success("Resume uploaded and read successfully")


client = Groq(api_key=os.environ.get("GROQ_API_KEY"))

if "rewritten_resume" not in st.session_state:
    st.session_state.rewritten_resume = ""


SECTION_NAMES = ["SUMMARY", "SKILLS", "EXPERIENCE", "EDUCATION"]


def clean_text(text):
    text = text.replace("—", "–")
    text = text.replace("●", "•")
    text = text.replace("* ", "• ")
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def clean_filename(name):
    name = name.strip()
    name = re.sub(r'[\\/*?:"<>|]', "", name)
    name = re.sub(r"\s+", "_", name)
    return name if name else "rewritten_resume"


def is_section(line):
    return line.strip().upper() in SECTION_NAMES


def is_bullet(line):
    stripped = line.strip()
    return stripped.startswith("•") or stripped.startswith("- ")


def clean_bullet(line):
    line = line.strip()
    if line.startswith("•"):
        return line[1:].strip()
    if line.startswith("-"):
        return line[1:].strip()
    return line


def looks_like_experience_header(line):
    line = line.strip()
    if "|" in line and "@" not in line and line.upper() not in SECTION_NAMES:
        return True
    return False


def add_bottom_border(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_run_font(run, size=9.2, bold=False):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold


def create_word_file(text):
    text = clean_text(text)
    lines = [line.strip() for line in text.split("\n") if line.strip()]

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.35)
    section.bottom_margin = Inches(0.35)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9.2)

    for i, line in enumerate(lines):
        line = line.strip()

        if i == 0:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(line)
            set_run_font(run, size=14.5, bold=True)

        elif i == 1:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(line)
            set_run_font(run, size=9.8, bold=True)

        elif i == 2:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(5)
            run = p.add_run(line)
            set_run_font(run, size=9.2, bold=False)

        elif is_section(line):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(line.upper())
            set_run_font(run, size=10.2, bold=True)
            add_bottom_border(p)

        elif is_bullet(line):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.first_line_indent = Inches(-0.18)
            p.paragraph_format.space_after = Pt(1.5)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(clean_bullet(line))
            set_run_font(run, size=9.1, bold=False)

        elif looks_like_experience_header(line):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(line)
            set_run_font(run, size=9.4, bold=True)

        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(line)
            set_run_font(run, size=9.1, bold=False)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def create_pdf_file(text):
    text = clean_text(text)
    lines = [line.strip() for line in text.split("\n") if line.strip()]

    buffer = io.BytesIO()

    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=0.45 * inch,
        leftMargin=0.45 * inch,
        topMargin=0.35 * inch,
        bottomMargin=0.35 * inch,
    )

    styles = getSampleStyleSheet()

    name_style = ParagraphStyle(
        "Name",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=14.5,
        leading=16,
        alignment=TA_CENTER,
        spaceAfter=1,
    )

    role_style = ParagraphStyle(
        "Role",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=9.8,
        leading=11,
        alignment=TA_CENTER,
        spaceAfter=1,
    )

    contact_style = ParagraphStyle(
        "Contact",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=9.2,
        leading=10.5,
        alignment=TA_CENTER,
        spaceAfter=5,
    )

    section_style = ParagraphStyle(
        "Section",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=10.2,
        leading=12,
        alignment=TA_LEFT,
        borderWidth=0.5,
        borderColor="black",
        borderPadding=1,
        spaceBefore=6,
        spaceAfter=3,
    )

    normal_style = ParagraphStyle(
        "NormalResume",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=9.1,
        leading=11.2,
        alignment=TA_LEFT,
        spaceAfter=2,
    )

    exp_header_style = ParagraphStyle(
        "ExperienceHeader",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=9.4,
        leading=11.2,
        alignment=TA_LEFT,
        spaceBefore=4,
        spaceAfter=1,
    )

    bullet_style = ParagraphStyle(
        "BulletResume",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=9.1,
        leading=11.2,
        leftIndent=13,
        firstLineIndent=-8,
        spaceAfter=1.5,
    )

    story = []

    for i, line in enumerate(lines):
        safe = html.escape(line)

        if i == 0:
            story.append(Paragraph(safe, name_style))

        elif i == 1:
            story.append(Paragraph(safe, role_style))

        elif i == 2:
            story.append(Paragraph(safe, contact_style))

        elif is_section(line):
            story.append(Paragraph(line.upper(), section_style))

        elif is_bullet(line):
            bullet_text = html.escape(clean_bullet(line))
            story.append(Paragraph("• " + bullet_text, bullet_style))

        elif looks_like_experience_header(line):
            story.append(Paragraph(safe, exp_header_style))

        else:
            story.append(Paragraph(safe, normal_style))

    doc.build(story)
    buffer.seek(0)
    return buffer


if st.button("Rewrite Resume"):
    if not resume_text or not job_description:
        st.warning("Please upload your resume and paste the job description.")
    else:
        rewrite_prompt = f"""
You are an expert ATS-focused resume rewriting engine.

Your job is to completely rewrite the candidate's resume so it aligns as strongly as possible with the target job description.

The rewritten resume must be tailored only to the job description provided by the user.
Do not assume any specific industry, role type, domain, tools, or keywords unless they are clearly present in the job description or logically supported by the original resume.

LOCKED INFORMATION — DO NOT CHANGE:
- Candidate name
- Email
- Phone number
- Company names
- College/university names
- Degree names

DO NOT INCLUDE:
- Projects section
- Project names
- Project bullets
- LinkedIn
- GitHub
- Certifications
- Additional sections
- Scores
- Notes
- Explanations
- Markdown tables

EVERYTHING ELSE CAN BE REWRITTEN:
- Role title under the candidate name
- Summary
- Skills
- Work experience job titles
- Work experience bullets
- Domain positioning
- Tools and keywords
- Resume storyline
- Responsibility framing
- Achievement framing

TARGET SETTINGS:
- Tailor intensity: Complete Rewrite
- Alignment level: Maximum alignment with the job description
- Resume length: Standard resume style
- Experience level: Junior to Mid unless the original resume clearly supports senior-level work
- Output must be clean, professional, ATS-friendly, and ready for PDF/Word export.

KEYWORD EXTRACTION RULE:
Before rewriting, silently extract the strongest keywords from the job description.

Extract from:
- required platforms and systems
- required technical tools
- programming or query languages
- BI and reporting tools
- data platforms, databases, data warehouses, and architecture terms
- domain-specific business terms
- must-have qualifications
- nice-to-have qualifications
- reporting responsibilities
- analytics responsibilities
- process improvement responsibilities

Then ensure the rewritten resume includes the strongest relevant keywords in:
- SKILLS
- SUMMARY
- EXPERIENCE bullets

Do not only use broad terms like data analysis, reporting, dashboards, or stakeholders if the job description gives more specific terms.

Use exact job description terms where appropriate, especially named platforms, systems, databases, cloud tools, BI tools, AI tools, and domain keywords.

ATS KEYWORD COVERAGE RULE:
The rewritten resume should include at least 80 percent of the job description's strongest relevant hard-skill keywords when they are reasonably connected to the candidate's background.

Prioritize:
1. Required platforms and systems
2. Required technical tools
3. Required programming or query languages
4. BI/reporting tools
5. Data architecture, database, and warehouse terms
6. Domain-specific business terms
7. Analytics methods and reporting concepts

A required JD keyword may appear in SKILLS if it is reasonably connected to the candidate's background or can be naturally supported in EXPERIENCE.

Do not invent certifications, employers, schools, degrees, or contact details.
Do not claim deep expert ownership of a tool unless the original resume supports it.
If a JD tool is important but not strongly supported, phrase it as exposure, reporting support, analysis support, dashboard support, data extraction, data validation, or related workflow experience where believable.

HEADER RULES:
The resume header must follow this exact structure:

Candidate Name
Target Role Title
Email | Phone Number

Rules:
- Use the candidate name exactly as provided in the original resume.
- On the second line, write the best matching target role title based on the job description.
- On the third line, include only email and phone number.
- Do not include LinkedIn.
- Do not include GitHub.
- Do not include location unless the job description specifically requires it.

OUTPUT STRUCTURE RULES:
The final resume must include only these sections, in this exact order:

Candidate Name
Target Role Title
Email | Phone Number

SUMMARY

SKILLS

EXPERIENCE

EDUCATION

Do not include any other sections.

FORMATTING RULES:
Each work experience must follow exactly this format:

Job Title | Company Name
Date
Location
• Bullet 1
• Bullet 2
• Bullet 3
• Bullet 4
• Bullet 5

Do NOT output long unbroken paragraphs under EXPERIENCE.

CRITICAL BULLET RULE STRICT:
Every bullet MUST begin with this exact symbol:
•

Every line under EXPERIENCE that describes work must start with:
•

Do NOT output:
- plain sentences under EXPERIENCE
- hyphen bullets
- missing bullets
- numbered bullets
- paragraphs under a job title

Each role MUST have EXACTLY 5 bullets.
No more, no less.

If any role has fewer than 5 bullets or more than 5 bullets:
Rewrite the entire EXPERIENCE section before returning the final output.

If any bullet does not start with •:
Rewrite the entire EXPERIENCE section before returning the final output.

ROLE AND DOMAIN ADAPTATION:
Infer the target role, domain, responsibilities, tools, skills, and keywords only from the job description.

Rewrite the resume so the candidate appears maximally aligned with that target role.

If the job description is for a different role than the original resume, reposition the candidate's existing experience toward that role while keeping the resume internally consistent and believable.

Do not force any specific domain.
Do not add role-specific terms unless the job description supports them.

SUMMARY RULES:
Write one short paragraph.

Rules:
- Maximum 3 lines only.
- Must be specific to the job description.
- Must include target role, target domain, strongest relevant tools or methods, and business impact.
- Must sound natural and direct.
- Must include 2–4 important hard-skill keywords from the job description.

Do NOT use generic phrases:
- delivers
- drives
- results-driven
- experienced
- skilled
- passionate
- highly motivated
- detail-oriented
- dynamic professional
- hardworking
- fast learner

If the summary sounds generic or could apply to any candidate:
Rewrite it before returning the final output.

SKILLS RULES:
Use the heading SKILLS.

Rules:
- Use ONLY 3 skill categories.
- Each category must be highly relevant to the job description.
- Each line must follow this format:
Category: Skill, Skill, Skill
- Focus only on technical and domain-relevant skills.
- Prioritize hard skills, named tools, platforms, systems, methods, KPIs, databases, BI tools, and role-specific concepts from the job description.
- Include the strongest searchable keywords from the job description when reasonably connected to the candidate's background.
- Keep relevant skills from the original resume.
- Remove unrelated skills.
- Do not include generic categories like Collaboration, Communication, Business Acumen, or Tools and Technologies.
- Do not create generic categories like Data Analysis unless the skills inside are specific to the job description.
- Do not list a required JD keyword only in SKILLS if it cannot be naturally supported anywhere in EXPERIENCE.
- A required JD keyword may appear in SKILLS if it can be reasonably connected to the candidate's background or naturally supported in EXPERIENCE.
- Make the skills section recruiter-searchable, specific, and ATS-friendly.

The SKILLS section must include:
- the strongest required platforms or systems from the JD
- the strongest required programming/query languages from the JD
- the strongest required BI/reporting tools from the JD
- the strongest required domain terms from the JD

EXPERIENCE RULES:
Use the heading EXPERIENCE.

Rewrite each work experience section to align with the target job description.

For each role:
- Use the locked company name exactly as provided.
- Keep dates if present.
- Keep location if present.
- Rewrite job title if needed to improve alignment.
- Use EXACTLY 5 bullets per role.
- Make responsibilities relevant to the job description.
- Add tools, methods, systems, platforms, databases, reporting concepts, and outcomes where appropriate.
- Make each role feel different.
- Keep the most recent role most aligned with the target job description.

Do not include unrelated work unless it can be reframed toward the target role.

Do not repeat the same type of work in every role.
Each role must highlight a different strength.

EXPERIENCE KEYWORD PLACEMENT RULE:
The most important job description keywords should appear naturally in the EXPERIENCE section, not only in SKILLS.

For required platforms/tools:
- Mention them in the most recent role if reasonably believable.
- If not directly supported by the original resume, use adjacent wording such as reporting workflows, dashboard support, data extraction, data validation, analytics support, or system reporting exposure.

For domain terms:
- Mention them in business context where natural.
- Do not overclaim domain expertise if the original resume does not support it.

For tools like BI platforms, databases, cloud systems, AI tools, Salesforce, ERP, CRM, financial platforms, or analytics platforms:
- Include them only where they fit the work described.
- If the resume does not prove deep ownership, phrase as exposure, support, reporting, dashboarding, analysis workflow, or data validation experience.

BULLET QUALITY RULES:
Each bullet must be specific, useful, and role-relevant.

Most bullets should include:
- Action
- Method, tool, system, or process
- Insight, decision, responsibility, or recommendation
- Business, operational, technical, customer, financial, compliance, or process impact

Every experience section must show:
analysis → action → outcome

Do not make every bullet sound the same.
Do not use vague filler.
Do not write empty statements.

BULLET RHYTHM RULES:
Each role must include bullet variety.

For every work experience:
- At least 2 bullets must NOT start with a direct action verb.
- At least 1 bullet should start with an impact, problem, insight, responsibility, or business context.
- Avoid making all bullets follow the same grammar pattern.

DECISION AND OWNERSHIP SIGNAL:
Where relevant to the target job description, include direct decision-oriented language.

Use wording such as:
- recommended
- guided
- prioritized
- evaluated
- improved
- translated findings into action
- supported leadership decisions
- identified gaps
- reduced friction
- improved visibility
- standardized processes
- strengthened reporting
- optimized workflows
- supported planning
- improved execution
- influenced prioritization
- shaped recommendations
- clarified tradeoffs

Only use wording that fits the target role.

Make the candidate sound like they influenced practical decisions, not just produced analysis.

VARIATION RULES:
Do not use the same sentence structure across bullets.
Do not reuse the same phrasing across roles.
Each role must use different verbs and language.

Avoid repeating these phrases too often:
- resulting in
- leading to
- using data to
- leveraged
- responsible for
- worked on
- helped with
- collaborated with stakeholders
- provided data-driven insights
- informed business decisions

Mix bullet styles:
- Some bullets can include metrics
- Some bullets can focus on tools or methods
- Some bullets can focus on business decisions
- Some bullets can focus on process improvement
- Some bullets can focus on stakeholder communication
- Some bullets can focus on technical execution

METRIC RULES:
Use quantifiable achievements where they make sense.

Rules:
- Include 1–2 quantified bullets in the most recent role.
- Include at least 1 quantified bullet in other roles when believable.
- Do not force a number into every bullet.
- Do not add fake-looking percentages everywhere.
- Do not repeat the same metric values across bullets.

Vary impact types:
- Percentage improvement
- Time saved
- Cost reduction
- Revenue influence
- Error reduction
- Process efficiency
- Reporting speed
- Customer/user impact
- Compliance improvement
- Quality improvement
- Operational improvement
- Decision speed
- Workload reduction

If exact numbers are not available, use realistic directional impact instead.

REALISM RULES:
The resume should align strongly with the job description, but it must still be internally consistent and believable.

Do not claim:
- Senior leadership ownership unless supported by the resume.
- Company-wide ownership unless supported by the resume.
- Tools that are completely unrelated to both the resume and job description.
- Responsibilities that are impossible for the candidate's experience level.
- Certifications, degrees, employers, or credentials not present in the original resume.

You may reframe experience, but do not invent new companies, schools, degrees, or contact details.

ROLE PROGRESSION RULES:
Do not make every role sound identical.

Earlier roles should sound more execution-focused:
- Support
- Reporting
- Documentation
- Cleaning
- Coordination
- Basic analysis
- Task ownership

Middle roles should show more ownership:
- Recurring reporting
- Analysis ownership
- Stakeholder communication
- Recommendations
- Process improvement
- KPI tracking
- Workflow support

Most recent role should show strongest alignment:
- Ownership
- Prioritization
- Decision support
- Cross-functional communication
- Optimization
- Business or technical recommendations
- Measurable impact

EDUCATION RULES:
Use the heading EDUCATION.

Keep education simple.
Do not rewrite university names.
Do not add new degrees.
Do not add coursework unless it directly supports the job description and fits naturally.

ATS FORMAT RULES:
Use a clean resume format.

Do not use:
- Markdown tables
- Decorative icons
- Columns
- Images
- Fancy formatting
- Explanations
- Comments
- Scores

Use only plain text section headings.

FINAL HUMANIZATION PASS:
Before final output, silently improve the resume so it does not sound machine-generated.

Check for:
- Repeated sentence patterns
- Too many bullets starting with the same type of verb
- Too many repeated phrases
- Generic summary opening
- Weak or generic skills section
- Missing JD-specific hard-skill keywords
- Missing named tools/platforms from the JD
- Weak decision-making language
- Missing measurable impact
- Same type of bullet repeated across multiple roles

Revise the resume internally before returning it.

FINAL ATS KEYWORD CHECK:
Before returning output, silently verify:
- The highest-value required tools/platforms from the JD appear in SKILLS if reasonably supportable.
- The highest-value technical skills from the JD appear in SKILLS.
- The highest-value BI/reporting/database/warehouse terms from the JD appear in SKILLS or EXPERIENCE.
- The highest-value domain terms from the JD appear where natural.
- Important keywords are not replaced only with vague terms.

If important hard-skill keywords from the JD are missing and can be reasonably supported:
Add them before returning.

FINAL VALIDATION MANDATORY:
Before returning output, verify:

1. Every role has EXACTLY 5 bullets.
2. Every bullet starts with •.
3. No plain sentences appear under EXPERIENCE.
4. Summary is maximum 3 lines and not generic.
5. Skills section has exactly 3 categories.
6. Skills section includes JD-specific hard skills, not only generic categories.
7. No repeated phrasing across roles.
8. No Projects section.
9. No LinkedIn or GitHub.

If any rule is violated:
Regenerate the output before returning.

FINAL CONSISTENCY CHECK:
Before producing the final resume, silently check:

- Is the resume aligned as strongly as possible to the job description?
- Did you use only these sections: SUMMARY, SKILLS, EXPERIENCE, EDUCATION?
- Did you preserve locked information?
- Are job titles believable?
- Are skills consistent with the rewritten experience?
- Does each role sound different?
- Are metrics varied and realistic?
- Is the summary specific to the target job?
- Is the output clean and ATS-friendly?
- Are there no contradictions?

FINAL OUTPUT RULES:
- Return only the rewritten resume.
- Do not explain what you changed.
- Do not include a score.
- Do not include notes.
- Do not include markdown tables.
- Do not include a Projects section.
- Do not include LinkedIn or GitHub.
- Use only these sections: SUMMARY, SKILLS, EXPERIENCE, EDUCATION.
- Fully rewrite the resume from top to bottom.
- Make the resume maximally aligned with the target job description.

OUTPUT FORMAT:

Candidate Name
Target Role Title
Email | Phone Number

SUMMARY

SKILLS

EXPERIENCE

EDUCATION

Original Resume:
{resume_text}

Target Job Description:
{job_description}
"""

        with st.spinner("Rewriting resume..."):
            response = client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[{"role": "user", "content": rewrite_prompt}],
                temperature=0.35
            )

        st.session_state.rewritten_resume = response.choices[0].message.content


if st.session_state.rewritten_resume:
    st.subheader("Edit Rewritten Resume Before Download")

    edited_resume = st.text_area(
        "You can edit the resume here",
        st.session_state.rewritten_resume,
        height=700
    )

    st.session_state.rewritten_resume = edited_resume

    file_name_input = st.text_input(
        "Enter file name before downloading",
        value="rewritten_resume"
    )

    safe_file_name = clean_filename(file_name_input)

    word_file = create_word_file(edited_resume)
    pdf_file = create_pdf_file(edited_resume)

    col1, col2 = st.columns(2)

    with col1:
        st.download_button(
            label="Download as Word",
            data=word_file,
            file_name=f"{safe_file_name}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    with col2:
        st.download_button(
            label="Download as PDF",
            data=pdf_file,
            file_name=f"{safe_file_name}.pdf",
            mime="application/pdf"
        )